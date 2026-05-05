"""DOCX document loader with support for images, tables, lists, and formatting."""

from pathlib import Path
from typing import Union, Optional, List, Dict, Any
import io

from app.services.document_loaders.base import (
    DocumentLoader, DocumentContent, ContentElement,
    ContentType, DocumentMetadata, ElementType
)
from app.core.logging import logger
from app.core.exceptions import DocumentProcessingError


class DOCXLoader(DocumentLoader):
    """DOCX loader with comprehensive content extraction."""
    
    supported_extensions = [".docx"]
    
    def load(self, file_path: Union[str, Path]) -> DocumentContent:
        """Load DOCX from file path."""
        file_path = Path(file_path)
        logger.info(f"Loading DOCX: {file_path}")
        
        try:
            with open(file_path, 'rb') as f:
                content = f.read()
            return self.load_from_bytes(content, file_path.name)
        except Exception as e:
            logger.error(f"Failed to load DOCX file: {str(e)}")
            raise DocumentProcessingError(f"DOCX load failed: {str(e)}")
    
    def load_from_bytes(self, content: bytes, filename: Optional[str] = None) -> DocumentContent:
        """Load DOCX from bytes."""
        try:
            from docx import Document
            
            doc = Document(io.BytesIO(content))
            elements = []
            
            # Extract metadata
            metadata = self._extract_metadata(doc, filename, len(content))
            
            # Process document elements
            current_list = []
            list_is_ordered = False
            in_list = False
            
            for para in doc.paragraphs:
                # Handle lists
                if self._is_list_item(para):
                    if not in_list:
                        in_list = True
                        list_is_ordered = self._is_ordered_list(para)
                        current_list = []
                    
                    text = self._extract_paragraph_text(para)
                    if text.strip():
                        current_list.append(text)
                else:
                    # Flush current list if any
                    if current_list:
                        elements.append(self._create_list_element(current_list, list_is_ordered))
                        current_list = []
                        in_list = False
                    
                    # Process paragraph
                    para_element = self._process_paragraph(para)
                    if para_element:
                        elements.append(para_element)
            
            # Flush remaining list
            if current_list:
                elements.append(self._create_list_element(current_list, list_is_ordered))
            
            # Extract tables
            table_elements = self._extract_tables(doc)
            elements.extend(table_elements)
            
            # Extract images
            image_elements = self._extract_images(doc)
            elements.extend(image_elements)
            
            # Get raw text
            raw_text = '\n\n'.join([e.to_text() for e in elements])
            
            # Create document content
            doc_content = DocumentContent(
                elements=elements,
                metadata=metadata,
                raw_text=raw_text
            )
            
            # Update word count
            metadata.word_count = len(raw_text.split())
            
            logger.info(f"DOCX loaded successfully: {len(elements)} elements")
            return doc_content
            
        except ImportError:
            logger.error("python-docx not installed")
            raise DocumentProcessingError("python-docx library required")
        except Exception as e:
            logger.error(f"DOCX processing failed: {str(e)}")
            raise DocumentProcessingError(f"DOCX processing failed: {str(e)}")
    
    def _extract_metadata(self, doc, filename: Optional[str], file_size: int) -> DocumentMetadata:
        """Extract DOCX metadata."""
        core_props = doc.core_properties
        
        return DocumentMetadata(
            title=core_props.title,
            author=core_props.author,
            subject=core_props.subject,
            keywords=core_props.keywords.split(',') if core_props.keywords else [],
            created_at=str(core_props.created) if core_props.created else None,
            modified_at=str(core_props.modified) if core_props.modified else None,
            source_file=filename,
            file_size=file_size,
            custom={
                "category": core_props.category,
                "comments": core_props.comments,
                "version": core_props.version,
            }
        )
    
    def _is_list_item(self, para) -> bool:
        """Check if paragraph is a list item."""
        if para._element.pPr is None:
            return False
        return para._element.pPr.numPr is not None
    
    def _is_ordered_list(self, para) -> bool:
        """Check if list item is part of ordered list."""
        try:
            numPr = para._element.pPr.numPr
            if numPr and numPr.numId:
                # Check numId to determine if it's numbered
                numId = numPr.numId.val
                return True  # Simplified - could check numbering definition
        except:
            pass
        return False
    
    def _extract_paragraph_text(self, para) -> str:
        """Extract text from paragraph with run formatting."""
        text_parts = []
        for run in para.runs:
            text = run.text
            if text:
                # Note formatting
                if run.bold:
                    text = f"**{text}**"
                if run.italic:
                    text = f"*{text}*"
                text_parts.append(text)
        
        return ''.join(text_parts)
    
    def _process_paragraph(self, para) -> Optional[ContentElement]:
        """Process a paragraph into a content element."""
        text = self._extract_paragraph_text(para)
        
        if not text.strip():
            return None
        
        # Check style for heading
        style_name = para.style.name if para.style else ""
        
        if style_name.startswith("Heading"):
            level = int(style_name.replace("Heading ", "")) if style_name.replace("Heading ", "").isdigit() else 1
            return self._create_heading_element(text, level)
        elif para.style and 'title' in style_name.lower():
            return self._create_heading_element(text, 1)
        elif para.style and 'subtitle' in style_name.lower():
            return self._create_heading_element(text, 2)
        elif self._is_code_block(para):
            return ContentElement(
                type=ContentType.CODE_BLOCK,
                content=text,
                metadata={"language": "", "style": style_name}
            )
        elif self._is_quote(para):
            return ContentElement(
                type=ContentType.QUOTE,
                content=text
            )
        else:
            return self._create_text_element(text, ContentType.PARAGRAPH)
    
    def _is_code_block(self, para) -> bool:
        """Detect code block by style."""
        style_name = para.style.name if para.style else ""
        code_styles = ['code', 'preformatted', 'html', 'xml', 'source']
        return any(s in style_name.lower() for s in code_styles)
    
    def _is_quote(self, para) -> bool:
        """Detect quote by style."""
        style_name = para.style.name if para.style else ""
        quote_styles = ['quote', 'quotation', 'intense quote']
        return any(s in style_name.lower() for s in quote_styles)
    
    def _extract_tables(self, doc) -> List[ContentElement]:
        """Extract tables from document."""
        elements = []
        
        for table_idx, table in enumerate(doc.tables, 1):
            rows = []
            headers = None
            
            for row_idx, row in enumerate(table.rows):
                cells = [cell.text.strip() for cell in row.cells]
                
                if row_idx == 0:
                    # First row as headers
                    headers = cells
                else:
                    rows.append(cells)
            
            if rows or headers:
                element = self._create_table_element(rows, headers)
                element.metadata.update({
                    "table_index": table_idx,
                    "row_count": len(rows),
                    "col_count": len(headers) if headers else 0,
                })
                elements.append(element)
        
        return elements
    
    def _extract_images(self, doc) -> List[ContentElement]:
        """Extract images from document."""
        elements = []
        image_count = 0
        
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                image_count += 1
                element = ContentElement(
                    type=ContentType.IMAGE,
                    content=f"[Image {image_count}]",
                    metadata={
                        "image_index": image_count,
                        "source": rel.target_ref,
                    }
                )
                elements.append(element)
        
        return elements
